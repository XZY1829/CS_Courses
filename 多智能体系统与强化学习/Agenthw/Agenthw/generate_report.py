from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# ============================================================
# Helper functions
# ============================================================
def set_run_font(run, name='宋体', size=Pt(12), bold=False, east_asia='宋体'):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)

def add_title(text, size=Pt(22)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, '黑体', size, bold=True, east_asia='黑体')
    return p

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(text, bold=False, indent=True, size=Pt(12)):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, '宋体', size, bold=bold)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return table


# ============================================================
# 封面
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
add_title('南京大学本科生课程设计报告', Pt(26))
doc.add_paragraph()
add_title('课程实验报告', Pt(22))
doc.add_paragraph()

info_lines = [
    '课程名称：多智能体与强化学习',
    '选择题目：峡谷漫步 / 重返秘境',
    '小组成员：（请填写姓名 学号）',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    set_run_font(run, '宋体', Pt(14))

doc.add_page_break()

# ############################################################
# 1  摘要
# ############################################################
add_heading_styled('1  摘要', level=1)

add_para(
    '本报告涵盖"多智能体与强化学习"课程的两次实验——"峡谷漫步"与"重返秘境"。'
    '两次实验均基于腾讯开悟（Kaiwu）强化学习平台，要求通过强化学习算法训练一个智能体，'
    '使其在包含障碍物的网格地图中学习移动策略，以最少的步数从起点导航到终点，并尽可能收集宝箱以获取更高分数。'
)
add_para(
    '实验一"峡谷漫步"（gorge_walk）的核心目标是在 64×64 的网格地图中训练智能体完成基本的路径导航任务。'
    '本实验采用 Dueling DQN 算法，将 Q 值函数分解为状态价值函数 V(s) 与优势函数 A(s,a)，'
    '通过共享特征层加双分支网络结构高效估计动作价值。训练过程中引入经验回放缓冲区和目标网络以增强稳定性，'
    '并结合课程学习、epsilon 指数衰减、多层次奖励塑造（包含距离引导、碰撞惩罚、振荡抑制和反停滞机制）等训练技巧。'
    '最终智能体以 64 步完成通关，获得 337 分。'
)
add_para(
    '实验二"重返秘境"（back_to_the_realm）在更大的 128×128 地图上进行，'
    '动作空间扩展为 8 方向移动与技能使用的组合（16 维），同时引入宝箱收集、增益道具和闪烁技能等新机制。'
    '本实验采用 Target-DQN 算法并加入梯度裁剪策略，使用 CNN+MLP 混合网络结构处理包含空间信息的复合观测'
    '（404 维向量特征与 4 通道 51×51 局部地图），配合合法动作掩码、BFS 路径距离计算以及多维度奖励塑造。'
    '最终智能体在 237 步内完成通关，收集 4 个宝箱获得 400 分宝箱奖励，使用 1 次闪烁技能并获取 1 个增益道具，综合得分 702 分。'
)
add_para(
    '两次实验的对比展示了从简单导航到复杂多目标任务的渐进式学习过程，体现了算法选型、'
    '网络架构设计、特征工程和奖励塑造等关键要素在强化学习应用中的重要性。'
)

# ############################################################
# 2  相关工作
# ############################################################
add_heading_styled('2  相关工作', level=1)

add_para(
    '深度强化学习（Deep Reinforcement Learning, DRL）将深度神经网络的表征学习能力与强化学习的决策优化框架相结合，'
    '近年来在游戏智能体、机器人控制、自动驾驶等领域取得了突破性进展。'
    '本节从值函数方法的演进、网格导航任务的研究以及奖励塑造技术三个维度梳理与本实验相关的工作。'
)

add_heading_styled('2.1 值函数方法的演进：从 DQN 到 Dueling DQN', level=2)
add_para(
    'Mnih 等人（2015）提出的深度 Q 网络（DQN）首次成功地将深度神经网络应用于高维状态空间的值函数逼近。'
    'DQN 的两项核心创新——经验回放（Experience Replay）和目标网络（Target Network）——'
    '分别通过打破样本时间相关性和稳定训练目标，解决了深度网络训练 Q 函数时的不稳定问题。'
    '在此基础上，研究者提出了多种改进方案。'
)
add_para(
    'Van Hasselt 等人（2016）提出 Double DQN，通过将动作选择与动作评估分离——使用在线网络选择最优动作、'
    '目标网络评估该动作的价值——有效缓解了标准 DQN 中的 Q 值过估计问题。'
    'Schaul 等人（2016）提出优先经验回放（Prioritized Experience Replay），'
    '根据 TD 误差对经验赋予不同的采样优先级，使得信息量更大的样本被更频繁地用于训练，显著提升了样本效率。'
)
add_para(
    'Wang 等人（2016）提出的 Dueling DQN 是本实验一采用的核心算法。其创新在于将 Q 值函数显式分解为'
    '状态价值 V(s) 和动作优势 A(s,a) 两个流，即 Q(s,a) = V(s) + A(s,a) - mean(A)。'
    '这种架构使网络能够独立学习哪些状态本身就有价值（无论采取何种动作），从而在动作对状态价值影响较小的场景下'
    '实现更高效的价值估计。Dueling DQN 在 Atari 游戏中的表现相比 DQN 有显著提升，'
    '尤其在动作空间较大但多数动作对结果影响相近的任务中表现突出。'
)

add_heading_styled('2.2 网格世界导航与路径规划', level=2)
add_para(
    '网格世界（Grid World）是强化学习领域最经典的实验环境之一。'
    '从 Sutton 和 Barto（2018）的教材中的简单迷宫到 MiniGrid（Chevalier-Boisvert 等, 2018）等复杂环境，'
    '网格导航任务一直是检验强化学习算法能力的重要基准。'
    '经典方法如 Q-learning、SARSA 等表格型方法可以解决小规模网格问题，'
    '但当状态空间增大（如本实验中的 64×64 和 128×128 网格）时，函数逼近方法成为必然选择。'
)
add_para(
    '在复杂网格导航中，局部观测和部分可观测性是常见挑战。'
    '研究者通常通过卷积神经网络（CNN）处理局部视野的空间信息，'
    '并结合循环网络（LSTM/GRU）或记忆机制来应对部分可观测问题。'
    '本实验二中采用的 CNN+MLP 混合架构正是遵循这一思路，'
    '使用 CNN 提取 51×51 局部视野中的障碍物、宝箱和终点等空间特征。'
    '此外，维护全局记忆图（memory map）记录访问历史也是一种常用的处理探索问题的手段，'
    '可视为一种简化的外部记忆机制。'
)

add_heading_styled('2.3 奖励塑造与课程学习', level=2)
add_para(
    '奖励塑造（Reward Shaping）是将领域知识注入强化学习过程的重要手段。'
    'Ng 等人（1999）从理论上证明了基于势函数的奖励塑造不会改变最优策略，为安全的奖励设计提供了理论基础。'
    '在实际应用中，距离引导奖励（鼓励靠近目标）、碰撞惩罚（抑制无效动作）、探索奖励（激励发现新区域）等'
    '是最常用的塑造策略。本实验中的多维度奖励设计综合运用了上述策略，'
    '尤其是实验二中按宝箱收集状态分阶段激活不同奖励项的方式，体现了层次化奖励设计的思想。'
)
add_para(
    '课程学习（Curriculum Learning, Bengio 等, 2009）通过由易到难地组织训练任务来加速收敛。'
    '在强化学习中，课程学习的典型实践包括逐步增大环境复杂度、分阶段引入任务目标等。'
    '本实验一中采用的"前 20% 轮次关闭宝箱、先学通关再学收集"的策略即为课程学习的直接应用，'
    '先让智能体掌握基本导航能力，再增加宝箱收集的附属任务。'
)
add_para(
    '此外，梯度裁剪（Gradient Clipping）也是深度强化学习训练中的重要稳定化技术。'
    'Pascanu 等人（2013）系统研究了循环网络中的梯度爆炸问题，提出通过限制梯度范数来稳定训练过程。'
    '本实验二中在 Target-DQN 的基础上引入 max_norm=10.0 的梯度裁剪，有效防止了 CNN 分支参数更新时可能出现的梯度爆炸。'
)

# ############################################################
# 3  实验环境及实验设置
# ############################################################
add_heading_styled('3  实验环境及实验设置', level=1)

add_heading_styled('3.1 实验平台', level=2)
add_para(
    '两次实验均在腾讯开悟（Kaiwu）强化学习平台上完成。开悟平台提供了基于浏览器的集成开发环境（IDE），'
    '支持在线代码编辑、训练任务管理、模型保存/加载、模型评估与回放等功能。'
    '实验采用 PyTorch 深度学习框架实现算法，通过平台提供的 kaiwu_agent SDK 与环境进行交互。'
)

add_heading_styled('3.2 实验一：峡谷漫步', level=2)

add_heading_styled('3.2.1 环境设置', level=3)
add_para('峡谷漫步环境（gorge_walk）的基本设置如下：')
add_table(
    ['属性', '描述'],
    [
        ['地图尺寸', '64×64 网格'],
        ['动作空间', '4 个离散动作（上、下、左、右）'],
        ['观测维度', '250 维特征向量'],
        ['终止条件', '到达终点（terminated）或超过最大步数（truncated）'],
    ]
)
doc.add_paragraph()

add_heading_styled('3.2.2 算法设计', level=3)
add_para(
    '本实验采用 Dueling DQN 算法。Dueling DQN 将 Q 值分解为状态价值 V(s) 和优势函数 A(s,a)：'
    'Q(s,a) = V(s) + A(s,a) - mean(A(s,·))。网络架构包括共享特征层（Linear(250,128)+ReLU）、'
    'Value 分支（Linear(128,128)+ReLU+Linear(128,1)）和 Advantage 分支（Linear(128,128)+ReLU+Linear(128,4)）。'
)
add_para(
    '状态特征包含 7 类信息：(1) 位置一维编码 pos_x×64+pos_z；'
    '(2) 行列 one-hot 编码（各 64 维）；(3) 到终点和宝箱的离散化距离；'
    '(4) 局部视野中的障碍物、宝箱、终点分布图（展平为向量）；'
    '(5) 局部视野中已访问区域的记忆图；(6) 宝箱收集状态。所有特征拼接后裁剪/填充至 250 维。'
)
add_para(
    '训练采用经验回放缓冲区（容量 10000）和目标网络（每 100 步硬更新），'
    '使用 Adam 优化器（学习率 1e-3），折扣因子 γ=0.9，批量大小 64。'
    'Epsilon 按 episode 指数衰减（0.9995^episode），从 1.0 衰减至 0.05。'
)
add_para(
    '奖励塑造包含多个维度：基础步数惩罚（-0.02/步）引导高效通关；'
    '距离引导（靠近终点 +0.08，远离 -0.05）提供方向信号；'
    '通关奖励/超时惩罚（±8.0）设定明确的成功/失败信号；'
    '循环抑制惩罚（原地 -0.15、振荡 -0.12、反向 -0.08）避免无效行为。'
    '此外引入课程学习策略——前 20% 轮次关闭宝箱生成，让智能体先学会基本导航；'
    '以及反停滞机制——连续 18 步无进展时强制随机探索 4 步。'
)

add_heading_styled('3.2.3 实验结果与分析', level=3)
add_table(
    ['指标', '值'],
    [
        ['得分', '337'],
        ['步数', '64'],
        ['宝箱收集数', '0'],
        ['宝箱得分', '0'],
    ]
)
doc.add_paragraph()
add_para(
    '智能体以 64 步完成通关，获得 337 分，表明 Dueling DQN 算法配合精心设计的奖励塑造能够有效学习峡谷漫步任务的导航策略。'
    '64 步的通关步数较为理想，说明智能体找到了一条较为高效的路径。'
    '宝箱收集数为 0，主要原因是：(1) 课程学习前期关闭了宝箱，智能体优先学习了导航能力；'
    '(2) 通关奖励（+8.0）远大于宝箱奖励，导致策略倾向于直奔终点。'
    '若要提升宝箱收集表现，可考虑加大宝箱奖励权重或延长课程学习中宝箱引入的过渡期。'
)

add_heading_styled('3.3 实验二：重返秘境', level=2)

add_heading_styled('3.3.1 环境设置', level=3)
add_para('重返秘境环境（back_to_the_realm）的基本设置如下：')
add_table(
    ['属性', '描述'],
    [
        ['地图尺寸', '128×128 网格（从 protobuf 格式的 map_info 动态构建）'],
        ['动作空间', '16 个离散动作（8 方向 × 2 技能选项），带合法动作掩码'],
        ['观测维度', '404 维向量 + 4 通道 51×51 局部特征图（共 4500 维输入）'],
        ['附加机制', '闪烁技能（Talent）、增益道具（Buff）、宝箱收集'],
        ['终止条件', '到达终点（terminated）或超过最大步数（truncated）'],
    ]
)
doc.add_paragraph()

add_heading_styled('3.3.2 算法设计', level=3)
add_para(
    '本实验采用 Target-DQN 算法并加入梯度裁剪（max_norm=10.0）以增强训练稳定性。'
    '网络采用 CNN+MLP 混合架构：CNN 分支由 3 层卷积层（4→16→32→64 通道，3×3 卷积核，BatchNorm+ReLU+MaxPool）'
    '组成，处理 4 通道 51×51 的局部特征图（障碍物图、终点图、宝箱图、记忆图），输出 4096 维特征；'
    'MLP 分支将 CNN 输出与 404 维向量特征拼接为 4500 维后，经三层全连接层（256→128→16）映射为各动作的 Q 值。'
    '网络权重使用 Kaiming 正态初始化。'
)
add_para(
    '向量特征（404 维）包含：归一化位置（2 维）、行列 one-hot 编码（256 维）、'
    '终点/宝箱的相对方向与距离（9+135 维）、Buff 和 Talent 可用性（2 维）。'
    'Preprocessor 类负责从 protobuf 原始观测中解析英雄位置、构建 128×128 网格、'
    '使用 BFS 计算到各目标的网格距离、维护全局记忆图和最近 100 步位置记录。'
    '算法在预测和训练时均利用环境提供的合法动作掩码过滤非法动作。'
)
add_para(
    '奖励塑造包含多维度加权信号：宝箱收集奖励（权重 1.0）为最高优先级；'
    '终点距离奖励（权重 0.5，在收集至少 1 个宝箱后激活）引导后期导航；'
    '宝箱距离奖励（权重 0.5）引导靠近最近宝箱；碰撞惩罚（权重 -1.0）抑制无效碰撞；'
    '探索奖励（权重 0.1）鼓励发现新区域、惩罚重复访问；'
    '超时惩罚（权重 -3.0）是相对于基线的改进，有效教导智能体珍惜时间。'
    '训练采用 Adam 优化器（学习率 1e-4），折扣因子 γ=0.95，'
    '目标网络每 500 步同步，epsilon 线性衰减至 0.1。'
)

add_heading_styled('3.3.3 实验结果与分析', level=3)
add_table(
    ['指标', '值'],
    [
        ['得分', '702'],
        ['步数', '237'],
        ['宝箱收集数', '4'],
        ['宝箱得分', '400'],
        ['技能使用次数', '1'],
        ['增益获取数', '1'],
    ]
)
doc.add_paragraph()
add_para(
    '智能体综合得分 702 分，在多个维度上表现均衡。'
    '(1) 宝箱收集：成功收集 4 个宝箱获得 400 分宝箱奖励，验证了以宝箱收集为优先的奖励设计策略的有效性。'
    '"终点奖励在收集至少 1 个宝箱后才激活"的设计成功引导智能体先收集宝箱再奔向终点。'
    '(2) 导航效率：237 步完成通关，在 128×128 的大地图上绕道收集 4 个宝箱后仍能到达终点，说明路径规划能力良好。'
    '(3) 技能与增益：使用 1 次闪烁技能和获取 1 个增益道具，表明智能体具备一定的技能利用能力。'
    '(4) 总体来看，CNN 分支对空间特征的提取、BFS 距离引导和分阶段奖励激活三者的结合是取得良好表现的关键。'
)

add_heading_styled('3.4 两次实验对比', level=2)
add_table(
    ['对比维度', '实验一：峡谷漫步', '实验二：重返秘境'],
    [
        ['算法', 'Dueling DQN + 经验回放', 'Target-DQN + 梯度裁剪'],
        ['网络结构', '全连接 MLP (250→128→Q)', 'CNN+MLP 混合 (图+向量→Q)'],
        ['动作空间', '4（上下左右）', '16（8 方向×2 技能）'],
        ['地图规模', '64×64', '128×128'],
        ['得分', '337', '702'],
        ['步数', '64', '237'],
        ['宝箱', '0 个 / 0 分', '4 个 / 400 分'],
    ]
)
doc.add_paragraph()

# ############################################################
# 4  总结和展望
# ############################################################
add_heading_styled('4  总结和展望', level=1)

add_heading_styled('4.1 总结', level=2)
add_para(
    '通过两次实验的完整实践，我们对深度强化学习在导航任务中的应用获得了系统性的认识。'
    '实验一中，Dueling DQN 配合经验回放和目标网络成功解决了 64×64 网格中的路径规划问题，'
    '课程学习和多层次奖励塑造有效加速了训练收敛。'
    '实验二中，面对更大的状态空间、更复杂的动作空间和多目标任务，'
    'Target-DQN 结合 CNN 空间特征提取、合法动作掩码和梯度裁剪等技术实现了较好的综合表现。'
)
add_para(
    '两次实验的核心经验可归纳为以下几点：'
    '第一，算法选型应匹配问题复杂度，简单任务适合轻量级网络（如纯 MLP），'
    '复杂任务需要更强的特征提取能力（如 CNN）；'
    '第二，奖励塑造是引导智能体行为的关键杠杆，奖励权重的设置直接决定了策略偏好'
    '（实验一中通关奖励主导导致忽视宝箱，实验二中宝箱奖励优先则成功引导了收集行为）；'
    '第三，训练稳定性技术（经验回放、目标网络、梯度裁剪）对于深度强化学习的成功训练不可或缺；'
    '第四，特征工程仍然重要——合理的状态表征（BFS 距离、记忆图、位置编码等）为智能体提供了充分的决策依据。'
)

add_heading_styled('4.2 不足', level=2)
add_para(
    '当前实验仍存在一些不足之处。'
    '实验一中智能体未能收集宝箱，反映出奖励设计在多目标权衡上的欠缺；'
    '实验二中虽然收集了 4 个宝箱，但仍有提升空间，可能与探索不够充分或路径规划不够全局化有关。'
    '此外，两次实验均未引入 Double DQN 的动作选择-评估分离机制，可能存在一定程度的 Q 值过估计；'
    '实验二未使用经验回放，样本利用效率有待提升。'
    '在网络架构方面，实验二的 CNN 分支较为简单，对大范围空间信息的捕捉能力有限。'
)

add_heading_styled('4.3 展望', level=2)
add_para(
    '未来的改进方向包括：'
    '(1) 算法层面，可以引入 Double DQN 减少过估计、Prioritized Experience Replay 提升样本效率、'
    'Noisy Network 替代 epsilon-greedy 实现更自适应的探索；'
    '(2) 网络架构层面，可以探索注意力机制（Attention）使智能体动态关注关键区域，'
    '或引入 LSTM/GRU 处理序列决策中的时序依赖；'
    '(3) 训练策略层面，可以采用更精细的课程学习方案（如自适应难度调节）、'
    '分层强化学习（将宝箱收集和终点导航分为高层子目标）、'
    '或结合模仿学习利用示范轨迹加速前期训练；'
    '(4) 奖励设计层面，可以探索逆向强化学习（Inverse RL）自动从专家行为中学习奖励函数，'
    '避免手工设计奖励权重的繁琐调参过程。'
)


# ============================================================
# 参考文献
# ============================================================
add_heading_styled('参考文献', level=1)
refs = [
    '[1] Mnih V, Kavukcuoglu K, Silver D, et al. Human-level control through deep reinforcement learning[J]. Nature, 2015, 518(7540): 529-533.',
    '[2] Van Hasselt H, Guez A, Silver D. Deep reinforcement learning with double Q-learning[C]. AAAI, 2016.',
    '[3] Wang Z, Schaul T, Hessel M, et al. Dueling network architectures for deep reinforcement learning[C]. ICML, 2016.',
    '[4] Schaul T, Quan J, Antonoglou I, et al. Prioritized experience replay[C]. ICLR, 2016.',
    '[5] Ng A Y, Harada D, Russell S. Policy invariance under reward transformations: Theory and application to reward shaping[C]. ICML, 1999.',
    '[6] Bengio Y, Louradour J, Collobert R, et al. Curriculum learning[C]. ICML, 2009.',
    '[7] Sutton R S, Barto A G. Reinforcement learning: An introduction[M]. MIT Press, 2018.',
    '[8] Pascanu R, Mikolov T, Bengio Y. On the difficulty of training recurrent neural networks[C]. ICML, 2013.',
]
for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    set_run_font(run, 'Times New Roman', Pt(10.5), east_asia='宋体')


# ============================================================
# Save
# ============================================================
output_path = r'c:\Users\zhengyang.xu\Desktop\WorkDoc\Study\Agenthw\实验报告.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
